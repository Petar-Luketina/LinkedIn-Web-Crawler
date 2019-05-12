# The various libraries need to be imported to make the job much quicker
# These contain useful pre-written code
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import math
import random
import datetime
from time import sleep
from docx import Document

# We must find the downloaded webdriver. We can use Chrome, FireFox, IE, etc.
# The webdriver will literally 'drive' much of the actions
driver = webdriver.Firefox(executable_path= "C:\\Users\\***CHANGE THIS PATH***\\Desktop\\LinkedIn-Web-Crawler\\geckodriver.exe")

# The specifications for the browser, which are unneeded but preferred.
profile = webdriver.FirefoxProfile()
profile.set_preference("browser.cache.disk.enable", False)
profile.set_preference("browser.cache.memory.enable", False)
profile.set_preference("browser.cache.offline.enable", False)
profile.set_preference("network.http.use-cache", False)
profile.update_preferences()

# The web browser finds LinkedIn's URL
driver.get(url = 'https://www.linkedin.com/')

# Sleep is used often to allow the browser to load
sleep(1)
# It also mimics a real person browsing the website




# The email login and password go here. We must hide it!
email = ''
password = ''

# The job we are looking for and where we want to look
desired_job = 'Python'
desired_location = 'Cleveland/Akron, Ohio Area'




# The locations to avoid
locations_to_avoid = ['Akron, OH, US',
                     'Kent, OH, US',
                     'Elyria, OH, US',
                     'Middlefield, OH, US',
                     'Westlake, OH, US',
                     'Fairlawn, OH, US', 'Fairlawn, Ohio, United States',
                     'Middleburg Heights, OH, US',
                     'Hudson, OH, US',
                     'Lorain, OH, US',
                     'Oberlin, OH, US',
                     'Strongsville, OH, US',
                     'Uniontown, OH, US',
                     'Wooster, OH, US',
                     'Sandusky, OH, US',
                     'Rittman, OH, US',
                     'North Olmsted, OH, US',
                    ]

# We can also avoid certain employers but I won't call out anyone here
employers_to_avoid = []

# These are the keywords that need to be avoided which vary for every position
keywords_to_avoid = ['c++', 'c ++', 'c + +',
                     'c#', 'c #',
                     '.net',
                     'senior', 'sr', 'sr.',
                     'hadoop',
                     'java',
                     'sas',
                     'manager', 'management',
                     'lead', 'leader',
                     'security',
                    ]




# We have to find the sign in button first
sign_in = driver.find_element(by='link text', value='Sign in')

# The click it, of course
sign_in.click()

# Let's let it load the next page
sleep(1)




# We find the inputs so the program can type the login credentials
email_entry = driver.find_element(by='css selector', value='#username')
password_entry = driver.find_element(by='css selector', value='#password')

# For every letter in our email, we will set a random time between keystrokes
for letter in email:
    sleep(random.uniform(.1, .4))
    email_entry.send_keys(letter)

# For every letter in our password, we will set a random time between keystrokes
for letter in password:
    sleep(random.uniform(.1, .4))
    password_entry.send_keys(letter)

# Now we hit ther 'return' key or the 'Enter' key
password_entry.send_keys(Keys.RETURN)
sleep(3)




# We find and click the search bar
search_bar = driver.find_element(by='css selector', value='input')
search_bar.click()
sleep(1)

# Then we find and click the "Jobs" button to start a job search
jobs = driver.find_element(by='link text', value='Jobs')
jobs.click()
sleep(2)



# Finds the input so we can type the job we want
role_input_element = '//input[@placeholder="Search jobs"]'
role = driver.find_element(by='xpath', value=role_input_element)

# Again, we cycle through each letter and wait between keystrokes
for letter in desired_job:
    sleep(random.uniform(.1, .35))
    role.send_keys(letter)

# We also have find the input for our desired location
location_input_element = '//input[@placeholder="Search location"]'
location = driver.find_element(by='xpath', value=location_input_element)

# Same deal af before... waiting between keystrokes as we type in our location
for letter in desired_location:
    sleep(random.uniform(.1, .35))
    location.send_keys(letter)

# Hit 'Return' and wait for the page to load
location.send_keys(Keys.RETURN)
sleep(2)




# Finds and clicks the 'Filter' button
filter_button_element = '//*/header/div/div/ul/li[5]/button'
filter_button = driver.find_element(by='xpath', value=filter_button_element)
filter_button.click()

# Now we choose which options to filter by

# Only posts from the past month... We don't want to waste time on old listings
past_month_element = '.ember-view ul > li:nth-child(3) > label > p'
date_posted = driver.find_element(by='css selector', value=past_month_element)
date_posted.click()

# Full time employement only, please    <($.$)<
full_time_element = '(//*[@class="ember-view"]/ul/li[1]/label/p/span[1])[4]'
job_type = driver.find_element(by='xpath', value=full_time_element)
job_type.click()
sleep(.5)

# I was late to the party so entry- and associate-level positions for me
entry_level_element = '(//*[@class="ember-view"]/ul/li[2]/label/p/span[1])[9]'
entry_level = driver.find_element(by='xpath', value=entry_level_element)
entry_level.click()
sleep(.5)
associate_level_element = '(//*[@class="ember-view"]/ul/li[3]/label/p/span[1])[9]'
associate_level = driver.find_element(by='xpath', value=associate_level_element)
associate_level.click()
sleep(.5)

# We submit our filters and wait for the filters to be applied
submit_element = '//artdeco-modal-header/div/div[2]/button[2]'
apply_filters = driver.find_element(by='xpath', value=submit_element)
apply_filters.click()
sleep(2)


# How will we sort the listings?
sort_by = driver.find_element(by='xpath', value='//*[@id="sort-by-select"]')
sort_by.click()
sleep(.5)

# Let's even sort the listings by date
by_date = driver.find_element(by='xpath', value='//*[@id="sort-by-select-options"]/li[2]')
by_date.click()
sleep(2)



# Here we find the XML path for each of the elements we will need.

# This is the element on the web browser we need most: the individual listings
# The {} indicate that we can put anything in there (but we'll stick to numbers)
job_element = '.jobs-search-results__list > li:nth-child({})'


# Here are the various elements inside of the listing...
# Like the title
title_element = './div/artdeco-entity-lockup/artdeco-entity-lockup-content/h3/a'
# And location
location_element = './div/artdeco-entity-lockup/artdeco-entity-lockup-content/h5'
# And who posted the listing
employer_element = './div/artdeco-entity-lockup/artdeco-entity-lockup-content/h4'


# Here we get the total number of results. We will use this later
x_results = 'div.jobs-search-two-pane__alerts > div > div'




# This keeps track of which page we are on
page_number = [1]

# A function used to cycle through all 25 listings on a page
def run_thru_listed_positions():

    # For every one of the 25 listings, run this code:
    for i in range(1, 26):

        # The last pages do not have 25 listings so we can't have the code fail
        try:
            # .format() allows us to add a number to the {} we saw above
            job = driver.find_element(by='css selector', value=job_element.format(i))

            # And every 3rd listing, we should scroll to load the listings below
            if i % 3 == 0:
                job.click()

            # Here we call the function to unpack each listing (see below)
            evalute_job_listing(job)

        # If it fails, it just means we are at the last page and nothing should happen
        except:
            pass

    # This lets us know the progress of our code
    print('Page ' + str(page_number[0]) + '/' + str(pages) + ' complete...')

    # We cannot forget to change what page we are on
    page_number[0] += 1

    # As long as the current page isn't above the page number of pages, it's OK
    # and we find and click to the subsequent pages
    if page_number[0] <= pages and page_number[0] != 40:
        next_page_element = '//*[@class="search-results-pagination-section"]/artdeco-pagination/ul/li[{}]'
        next_page = driver.find_element(by='xpath', value=next_page_element.format(page_number[0]))
        next_page.click()
        sleep(2)




# The list where we will house all our desirable positions
found_desirable_positions = []

# An aptly named function that judges each listing
def evalute_job_listing(job):

    # We are targeting the title, location and employer of each listing
    title = job.find_element(by='xpath', value=title_element)
    location = job.find_element(by='xpath', value=location_element)
    employer = job.find_element(by='xpath', value=employer_element)

    # If the location isn't too far, we are in business.
    if location.text not in locations_to_avoid:
        status = True

        try:

            # If the employer checks out then its status is still good
            if employer.text not in employers_to_avoid:
                status = True

            else:
                return False

        except:

            # Any listing without an employer is avoided
            return False

    # Location is too far
    else:
        return False

    # If both the location and employer are okay, we run this code:
    if status == True:

        # The listings' titles need to be broken up
        for keyword in (title.text).split():

            # Each word is compared against the keywords to avoid
            if keyword.lower() not in keywords_to_avoid:
                status = True

            else:
                return False

    # If all the paramaters check out, we run the code below
    if status == True:

        # We do not want duplicates
        if job.text not in found_desirable_positions:

            # Let's add the URL to each listing
            url = '\n' + title.get_attribute('href')

            # We add all the listings with their URLS to a list
            found_desirable_positions.append(job.text + url)




# Instantiating a class from a library to use later
document = Document()

# Gets the current date and time so we can keep track when the code was finished
currentDT = datetime.datetime.now()

# The date-time must be formatted nicely (e.g. Firday, May 10, 2019)
formatted_time = currentDT.strftime("%a, %b %d, %Y")

# We title the document appropriately
document.add_heading('LinkedIn Jobs for ' + str(formatted_time), 0)

# We get the element that returns the total results from our LinkedIn search
# Then we turn "194 results" into "194" by cutting out the last eight letters
# The "194" must also be turned into an integer so we can devide with it
results = int(driver.find_element(by='css selector', value=x_results).text[0:-8])
pages = math.ceil(results/25)

# For every page, we will call the function that cycles through the listings
for i in range(0, pages):
    run_thru_listed_positions()

# We will be left with a list of desirable positions
# We will cycle through those listings and add it to our ducument
for listing in found_desirable_positions:
    print('\n' + listing)
    document.add_paragraph(('\n' + listing), style='List Number')

# The last step is naming and saving our document
document.save('LinkedIn Jobs - ' + str(formatted_time) + '.docx')
